import time

import docx2pdf
import pandas as pd

# 파이썬으로 엑셀파일 만들기
df = pd.DataFrame([["홍길동", "1990.01.02", '2021-0001'],
                   ["김민준", "1990.05.06", '2021-0002'],
                   ["김철수", "2000.08.08", '2021-0003'],
                   ["김영희", "2000.09.09", '2021-0004'],
                   ["이서준", "2010.10.10", '2021-0005'],
                   ["장다인", "2017.12.12", '2021-0006']])

print(df)
df.to_excel(r'수료증명단.xlsx', index=False, header=False)


# 엑셀의 정보를 불러와 수료증 자동 생성
from openpyxl import load_workbook

load_wb = load_workbook(r"수료증명단.xlsx")
load_ws = load_wb.active

name_list = []
birthday = []
ho_list = []
for i in range(1,load_ws.max_row + 1):
    name_list.append(load_ws.cell(i,1).value)
    birthday.append(load_ws.cell(i,2).value)
    ho_list.append(load_ws.cell(i,3).value)

print(name_list)
print(birthday)
print(ho_list)


# 수료증 내용을 채운 후 저장하는 코드 만들기

import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


doc = docx.Document(r'수료증양식.docx')

style = doc.styles['Normal']
style.font.name = '나눔고딕'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
style.font.size = docx.shared.Pt(12)

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run(' 제 2020-0001 호 \n')
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run(' 수 료 증 \n')
run.font.name = '나눔고딕'
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(40)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run(' 성 명: 장다인\n')
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
run = para.add_run(' 생 년 월 일: 2017.12.12\n')
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
run = para.add_run(' 교 육 과 정: 파이썬과 40개의 작품들\n')
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)
run = para.add_run(' 교 육 날 짜: 2021.08.05~2021.09.09\n')
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
run.font.size = docx.shared.Pt(20)

# 책에는 더 있지만 추가할 의미를 느낄수 없어 여기까지 적음
doc.save(r'수료증결과.docx')


# 처음 만들 엑셀을 이용해서 워드 만들고 pdf로 변환하기
from docx import Document
from openpyxl import load_workbook
import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert

for i in range(len(name_list)):
    doc = docx.Document(r'수료증양식.docx')

    style = doc.styles['Normal']
    style.font.name = '나눔고딕'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    style.font.size = docx.shared.Pt(12)

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run(' 제 ' +ho_list[i] +  '호 \n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run(' 수 료 증 \n')
    run.font.name = '나눔고딕'
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(40)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    run = para.add_run('\n\n')
    run = para.add_run(' 성 명: ' + name_list[i] + '\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run(' 생 년 월 일: ' + birthday[i] + '\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run(' 교 육 과 정: 파이썬과 40개의 작품들\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)
    run = para.add_run(' 교 육 날 짜: 2021.08.05~2021.09.09\n')
    run.font.name = '나눔고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '나눔고딕')
    run.font.size = docx.shared.Pt(20)

    a = doc.save(name_list[i] + '.docx')
    docx2pdf.convert(name_list[i] + '.docx',name_list[i] +str(i)+'.pdf')